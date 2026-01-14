#!/usr/bin/env python3
"""
ingest_sources_cli.py — optimized + source-identity aware

Features added:
- Memory-sane batching; pipeline lifecycle per batch by default
- Table preservation options (policy, rows-per-chunk, fallbacks)
- Strict source identity extraction per document (PDF/DOCX/HTML/TXT) with evidence & calibrated confidence
- Citation propagation directives so chunks include source meta (to be honored by the pipeline)

Note: This script expects your project modules under "app/". If imports fail, ensure this file is located
inside the repository and "app" is at or above its parent directory.
"""

import argparse
import asyncio
import json
import os
import re
import sys
import hashlib
import time
import gc
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple, Any

# --- Project path bootstrap (robust) ---
_here = Path(__file__).resolve().parent
# Extend search paths so we can resolve the real app package when the repo
# stores it under a service directory (e.g., "rag-service/app").
_candidates = [
    _here,
    _here.parent,
    _here.parent.parent,
    _here / "rag-service",
    _here.parent / "rag-service",
]

def _add_site_packages(base: Path) -> None:
    """Include virtualenv site-packages if present."""
    if not base.exists():
        return
    for site_root in ("lib", "lib64"):
        for site_path in (base / site_root).glob("python*/site-packages"):
            if site_path.exists() and str(site_path) not in sys.path:
                sys.path.insert(0, str(site_path))

for _cand in _candidates:
    if (_cand / "app").exists():
        if str(_cand) not in sys.path:
            sys.path.insert(0, str(_cand))
        _add_site_packages(_cand / "venv")
        continue
    # If this candidate corresponds to a parent of rag-service, still search for its venv
    if (_cand / "rag-service" / "venv").exists():
        _add_site_packages(_cand / "rag-service" / "venv")

# --- Project imports (assumed to exist in your repo) ---
try:
    from app.core.vectorstore import VectorStoreManager
    from app.services.cache import CacheService
    from app.pipelines.ingestion import IngestionPipeline
    from app.core.config import settings
    # If your project defines a DocumentType enum, import it; else fall back to a local light enum
    from app.models.documents import (
        DocumentType as _DocType,
        DocumentIngestionRequest,
        DocumentIngestionResponse,
    )
    DocumentType = _DocType
except (ModuleNotFoundError, ImportError):
    # Fallback minimal types so this script remains syntactically valid for review;
    # in your real environment, the "app" imports should resolve.
    class DocumentType:
        PDF = type("T", (), {"value": "pdf"})
        TEXT = type("T", (), {"value": "text"})
        MARKDOWN = type("T", (), {"value": "markdown"})
        WEB = type("T", (), {"value": "web"})
        DOCX = type("T", (), {"value": "docx"})
        CSV = type("T", (), {"value": "csv"})
        XLSX = type("T", (), {"value": "xlsx"})
    VectorStoreManager = object  # type: ignore
    CacheService = object  # type: ignore
    IngestionPipeline = object  # type: ignore
    class DocumentIngestionRequest(dict):  # type: ignore
        def __init__(self, **kwargs):
            super().__init__(**kwargs)

    class DocumentIngestionResponse(dict):  # type: ignore
        def __init__(self, **kwargs):
            super().__init__(**kwargs)

    class _FakeSettings:
        log_level = "INFO"
        openai_api_key = None
        google_api_key = None
    settings = _FakeSettings()

# --- Constants & defaults ---

if hasattr(settings, "chroma_persist_directory"):
    _chroma_candidates = []
    # Prefer the repo's rag-service/chroma_db if it exists
    _chroma_candidates.append((_here / "rag-service" / "chroma_db").resolve())
    current_chroma = Path(settings.chroma_persist_directory)
    if current_chroma.is_absolute():
        _chroma_candidates.append(current_chroma)
    else:
        _chroma_candidates.append((_here / current_chroma).resolve())
    for _chroma_dir in _chroma_candidates:
        if _chroma_dir.exists() or _chroma_dir.parent.exists():
            settings.chroma_persist_directory = str(_chroma_dir)
            try:
                _chroma_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
            break

DOCUMENT_TYPE_BY_EXTENSION: Dict[str, Any] = {
    ".pdf": DocumentType.PDF,
    ".txt": DocumentType.TEXT,
    ".md": getattr(DocumentType, "MARKDOWN", DocumentType.TEXT),
    ".markdown": getattr(DocumentType, "MARKDOWN", DocumentType.TEXT),
    ".html": getattr(DocumentType, "WEB", DocumentType.TEXT),
    ".htm": getattr(DocumentType, "WEB", DocumentType.TEXT),
    ".doc": getattr(DocumentType, "DOCX", DocumentType.TEXT),    # will be converted/handled upstream
    ".docx": getattr(DocumentType, "DOCX", DocumentType.TEXT),
    ".csv": getattr(DocumentType, "CSV", DocumentType.TEXT),
    ".xls": getattr(DocumentType, "XLSX", DocumentType.TEXT),
    ".xlsx": getattr(DocumentType, "XLSX", DocumentType.TEXT),
}

DEFAULT_PATTERNS = (
    "**/*.pdf", "**/*.txt", "**/*.html", "**/*.htm",
    "**/*.doc", "**/*.docx"
)

STATE_FILENAME = ".ingest_state.json"

DEFAULT_ENV_PATHS: List[Path] = [
    Path(".env"),
    Path("config.env"),
    Path.home() / ".config" / "rag" / "env",
    Path("/etc/profile.d/rag_env.sh"),
]

# --- Utilities ---

def iter_files(patterns: Sequence[str], base_dir: Path) -> List[Path]:
    seen: Dict[str, Path] = {}
    for pat in patterns:
        pat_path = (base_dir / pat) if not pat.startswith("/") else Path(pat)
        for p in pat_path.parent.glob(pat_path.name):
            if p.is_file():
                seen[str(p.resolve())] = p.resolve()
    return sorted(seen.values())

def infer_document_type(path: Path) -> Optional[Any]:
    return DOCUMENT_TYPE_BY_EXTENSION.get(path.suffix.lower())

def human_size(num_bytes: int) -> str:
    units = ["B", "KB", "MB", "GB", "TB"]
    size = float(num_bytes)
    for unit in units:
        if size < 1024.0 or unit == units[-1]:
            return f"{size:.1f}{unit}"
        size /= 1024.0

def make_title(path: Path) -> str:
    name = path.stem.replace("_", " ").replace("-", " ").strip()
    return re.sub(r"\s+", " ", name).title()


# --- Canonical Title Mapping ---

_TITLE_CONFIG = None

def load_title_config() -> Dict[str, Any]:
    """Load canonical title configuration from JSON file."""
    global _TITLE_CONFIG
    if _TITLE_CONFIG is not None:
        return _TITLE_CONFIG

    config_path = Path(__file__).parent / "rag-service" / "config" / "document_titles.json"
    if not config_path.exists():
        _TITLE_CONFIG = {"explicit_mappings": {}, "fam_chapters": {}, "pattern_rules": [], "fallback": {}}
        return _TITLE_CONFIG

    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            _TITLE_CONFIG = json.load(f)
    except Exception:
        _TITLE_CONFIG = {"explicit_mappings": {}, "fam_chapters": {}, "pattern_rules": [], "fallback": {}}

    return _TITLE_CONFIG


def resolve_canonical_title(path: Path) -> Optional[str]:
    """
    Resolve filename to canonical document title using configuration.
    Returns None if no match found, then fallback to make_title().
    """
    config = load_title_config()
    filename = path.name

    # 1. Check explicit mappings first
    explicit = config.get("explicit_mappings", {})
    if filename in explicit:
        return explicit[filename]

    # 2. Try pattern rules with FAM chapter lookups
    fam_chapters = config.get("fam_chapters", {})
    pattern_rules = config.get("pattern_rules", [])

    for rule in pattern_rules:
        regex = rule.get("regex")
        if not regex:
            continue

        match = re.match(regex, filename)
        if not match:
            continue

        # Build template variables
        groups = match.groups()
        template = rule.get("template", "")

        # Handle FAM chapter pattern
        if "{chapter}" in template:
            chapter = groups[0] if len(groups) > 0 else ""
            section = groups[1] if len(groups) > 1 else ""
            subsection = groups[2] if len(groups) > 2 else ""

            # Build description key
            desc_key_template = rule.get("description_key", "{chapter}")
            desc_key = desc_key_template.format(chapter=chapter, section=section, subsection=subsection)

            # Lookup description
            description = fam_chapters.get(desc_key)

            # Build title
            title = template.format(chapter=chapter, section=section, subsection=subsection, title="")

            # Add description if found
            if description:
                title = f"{title}: {description}"

            return title

        # Handle other templates
        elif "{title}" in template:
            # Extract title from capture group and transform
            title_value = groups[2] if len(groups) > 2 else ""
            if rule.get("transform") == "title_case":
                title_value = title_value.replace("-", " ").replace("_", " ").title()

            title = template.format(
                chapter=groups[0] if len(groups) > 0 else "",
                section=groups[1] if len(groups) > 1 else "",
                title=title_value
            )
            return title

    # 3. No match found
    return None


def load_env_file(path: Path) -> None:
    text = path.read_text(encoding="utf-8", errors="ignore")
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("export "):
            line = line[len("export "):]
        m = re.match(r'^([A-Za-z_][A-Za-z0-9_]*)=(.*)$', line)
        if not m:
            continue
        k, v = m.group(1), m.group(2).strip().strip('"').strip("'")
        os.environ.setdefault(k, v)

def compute_file_signature(path: Path) -> Dict[str, Any]:
    h = hashlib.md5()
    size = path.stat().st_size
    mtime = int(path.stat().st_mtime)
    with path.open("rb") as f:
        while True:
            chunk = f.read(1024 * 1024)
            if not chunk: break
            h.update(chunk)
    return {"hash": h.hexdigest(), "size": size, "mtime": mtime}

def load_ingest_state(base_dir: Path) -> Dict[str, Any]:
    p = base_dir / STATE_FILENAME
    if not p.exists():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8", errors="ignore"))
    except Exception:
        return {}

def save_ingest_state(base_dir: Path, state: Dict[str, Any]) -> None:
    p = base_dir / STATE_FILENAME
    tmp = p.with_suffix(".tmp")
    tmp.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(p)

# --- Source identity extraction per contract ---

@dataclass
class Evidence:
    cue: str
    page: Optional[int]
    section: str
    start: Optional[int] = None
    end: Optional[int] = None
    why: str = ""

def _calibrate_confidence(kind: str, signals: Dict[str, Any]) -> float:
    """
    kind: 'pdf'|'docx'|'html'|'txt'
    signals: flags like has_published_by, has_copyright, has_masthead, has_canonical, metadata_publisher_match
    """
    score = 0.0
    if signals.get("metadata_publisher_match"):
        score += 0.45
    if signals.get("has_published_by"):
        score += 0.45
    if signals.get("has_masthead"):
        score += 0.35
    if signals.get("has_canonical"):
        score += 0.15
    if signals.get("has_footer_repeat"):
        score += 0.20
    # Cap between 0 and 1
    score = max(0.0, min(1.0, score))
    # Nudge by type
    if score >= 0.95:
        return 0.97
    if score >= 0.85:
        return 0.90
    if score >= 0.5:
        return 0.70
    return score

def _classify_org_type(name: str) -> str:
    n = name.lower()
    if any(w in n for w in ["department", "ministry", "secretariat", "treasury", "agency", "commission"]):
        return "department"
    # Heuristic: default to organization; caller can override
    return "organization"

def extract_source_identity(path: Path, doc_type: Optional[Any]) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """
    Returns (source, document) per contract. Never infers from filename; uses only content/metadata.
    When insufficient, returns source.name='unknown' with confidence 0.0 and candidates[].
    """
    ext = path.suffix.lower()
    source = {
        "name": "unknown",
        "type": "unknown",
        "confidence": 0.0,
        "evidence": [],
        "candidates": []
    }
    document = {
        "title": None,
        "published_date": None,
        "canonical_url": None,
        "identifiers": {"doi": None, "isbn": None, "gov_id": None}
    }

    # Helper to add evidence
    def add_evidence(cue: str, section: str, page: Optional[int] = None, start: Optional[int] = None, end: Optional[int] = None, why: str = ""):
        source["evidence"].append({
            "cue": cue,
            "location": {"page": page, "section": section, "start": start, "end": end},
            "why": why or "Decisive cue per extraction priority"
        })

    # HTML
    if ext in (".html", ".htm"):
        try:
            data = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return source, document
        # Head tags
        og_site = re.search(r'<meta[^>]+property=["\']og:site_name["\'][^>]+content=["\']([^"\']+)["\']', data, re.I)
        canonical = re.search(r'<link[^>]+rel=["\']canonical["\'][^>]+href=["\']([^"\']+)["\']', data, re.I)
        meta_pub = re.search(r'<meta[^>]+name=["\']publisher["\'][^>]+content=["\']([^"\']+)["\']', data, re.I)
        title_tag = re.search(r'<title[^>]*>(.*?)</title>', data, re.I|re.S)
        og_url = re.search(r'<meta[^>]+property=["\']og:url["\'][^>]+content=["\']([^"\']+)["\']', data, re.I)

        site_name = None
        if og_site:
            site_name = og_site.group(1).strip()
            add_evidence(f'og:site_name={site_name}', section="html_head", page=None, start=og_site.start(1), end=og_site.end(1), why="Strong site masthead signal")
        if meta_pub:
            pub = meta_pub.group(1).strip()
            add_evidence(f'meta[name=publisher] content={pub}', section="html_head", start=meta_pub.start(1), end=meta_pub.end(1), why="Explicit publisher meta")
            # prefer explicit publisher if present
            site_name = site_name or pub
        if canonical:
            url = canonical.group(1).strip()
            document["canonical_url"] = url
            add_evidence(f'rel=canonical href={url}', section="html_head", start=canonical.start(1), end=canonical.end(1), why="Canonical URL for validation")
        if og_url and not document["canonical_url"]:
            document["canonical_url"] = og_url.group(1).strip()
            add_evidence(f'og:url={document["canonical_url"]}', section="html_head", start=og_url.start(1), end=og_url.end(1), why="og:url as canonical fallback")
        if title_tag:
            title_text = re.sub(r'\s+', ' ', title_tag.group(1)).strip()
            if title_text:
                document["title"] = title_text

        signals = {
            "metadata_publisher_match": bool(meta_pub),
            "has_published_by": False,
            "has_masthead": bool(site_name),
            "has_canonical": bool(document["canonical_url"]),
            "has_footer_repeat": False,
        }
        if site_name:
            source["name"] = site_name
            source["type"] = _classify_org_type(site_name) if site_name else "website"
            source["confidence"] = _calibrate_confidence("html", signals)
            if source["confidence"] < 0.85:
                # candidate set
                source["candidates"] = [{
                    "name": site_name,
                    "type": source["type"],
                    "confidence": 0.7,
                    "why": "Site head provided site_name but no explicit ownership phrasing"
                }]
        return source, document

    # PDF (PyPDF2 best effort)
    if ext == ".pdf":
        try:
            import PyPDF2  # type: ignore
            reader = PyPDF2.PdfReader(str(path))
            meta = getattr(reader, "metadata", None) or getattr(reader, "documentInfo", None)
            meta_dict = {}
            if meta:
                for k in dir(meta):
                    if k.startswith("/"):
                        try:
                            meta_dict[k] = getattr(meta, k)
                        except Exception:
                            pass
                    # PyPDF2 >=3 uses attributes
                for k in ("author", "creator", "producer", "title"):
                    v = getattr(meta, k, None)
                    if v: meta_dict[k] = v
            title = meta_dict.get("/Title") or meta_dict.get("title")
            author = meta_dict.get("/Author") or meta_dict.get("author")
            producer = meta_dict.get("/Producer") or meta_dict.get("producer")
            creator = meta_dict.get("/Creator") or meta_dict.get("creator")

            if title and not document["title"]:
                document["title"] = str(title)

            # Extract first page text to search for publisher cues
            first_text = ""
            page0 = None
            try:
                if reader.pages:
                    page0 = 1
                    first_text = reader.pages[0].extract_text() or ""
            except Exception:
                first_text = ""

            # Look for explicit statements
            pub_line = None
            if first_text:
                m = re.search(r'(?i)(Published|Issued)\s+by\s*:\s*(.+)', first_text)
                if m:
                    pub_line = m.group(2).strip()
                    add_evidence(f"{m.group(0)}", section="title", page=1, why="Formal front-matter statement")
                c = re.search(r'(?i)©\s+([^\n]+)', first_text)
                if c:
                    c_org = c.group(1).strip()
                    add_evidence(f"© {c_org}", section="title", page=1, why="Copyright holder as ownership cue")
                    # candidate
                    if not pub_line:
                        pub_line = c_org

            # Prefer explicit pub_line; else try author/company-like fields
            candidates = []
            if pub_line:
                candidates.append(("publisher", pub_line))
            for tag, val in (("author", author), ("creator", creator), ("producer", producer)):
                if val and isinstance(val, str) and len(val.strip()) >= 2:
                    candidates.append((tag, val.strip()))

            chosen = None
            for tag, val in candidates:
                # Drop obviously tool producers (e.g., "Microsoft Word", "PDFKit")
                if re.search(r"adobe|pdf|microsoft|word|excel|libreoffice|pdfkit|ghostscript|printer|distiller", val, re.I):
                    continue
                chosen = val
                break

            signals = {
                "metadata_publisher_match": bool(chosen and (chosen == author or chosen == creator)),
                "has_published_by": bool(pub_line),
                "has_masthead": False,
                "has_canonical": False,
                "has_footer_repeat": False,
            }
            if chosen:
                source["name"] = chosen
                source["type"] = _classify_org_type(chosen)
                source["confidence"] = _calibrate_confidence("pdf", signals)
                # Evidence for metadata fields
                if author:
                    add_evidence(f"PDF metadata Author={author}", section="metadata", page=None, why="PDF metadata field")
                if creator:
                    add_evidence(f"PDF metadata Creator={creator}", section="metadata", page=None, why="PDF metadata field")
                if producer:
                    add_evidence(f"PDF metadata Producer={producer}", section="metadata", page=None, why="PDF metadata field")
                if title:
                    add_evidence(f"PDF metadata Title={title}", section="metadata", page=None, why="Document title")
            else:
                # Insufficient signals
                source["candidates"] = []
                if author:
                    source["candidates"].append({"name": author, "type": "individual", "confidence": 0.4, "why": "Author metadata present but not a publisher"})
                if creator:
                    source["candidates"].append({"name": creator, "type": "organization", "confidence": 0.35, "why": "Creator tool/field present"})
                if producer:
                    source["candidates"].append({"name": producer, "type": "organization", "confidence": 0.25, "why": "Producer field often a tool"})
            return source, document
        except Exception:
            # Fallthrough to unknown; do not guess
            return source, document

    # DOC/DOCX (python-docx)
    if ext in (".docx", ".doc"):
        try:
            import docx  # type: ignore
            d = docx.Document(str(path))
            props = getattr(d, "core_properties", None)
            company = getattr(props, "company", None) if props else None
            author = getattr(props, "author", None) if props else None
            title = getattr(props, "title", None) if props else None
            if title:
                document["title"] = title
            # Scan first ~20 paragraphs for ownership cues
            text_first = []
            for p in d.paragraphs[:20]:
                t = p.text.strip()
                if t:
                    text_first.append(t)
            joined = "\n".join(text_first)
            pub_line = None
            m = re.search(r'(?i)(Published|Issued)\s+by\s*:\s*(.+)', joined)
            if m:
                pub_line = m.group(2).strip()
                add_evidence(m.group(0), section="frontmatter", page=1, why="Formal front-matter")
            c = re.search(r'(?i)©\s+([^\n]+)', joined)
            if c:
                c_org = c.group(1).strip()
                add_evidence(f"© {c_org}", section="frontmatter", page=1, why="Copyright holder as cue")
                if not pub_line:
                    pub_line = c_org
            candidates = []
            if pub_line:
                candidates.append(("publisher", pub_line))
            if company:
                candidates.append(("company", company))
                add_evidence(f"DOCX core Company={company}", section="metadata", why="DOCX property")
            if author:
                candidates.append(("author", author))
                add_evidence(f"DOCX core Author={author}", section="metadata", why="DOCX property")
            chosen = None
            for tag, val in candidates:
                if val and isinstance(val, str) and len(val.strip()) >= 2:
                    # Skip obvious tool strings
                    if re.search(r"microsoft|pdf|printer|converter", val, re.I):
                        continue
                    chosen = val.strip()
                    break
            signals = {
                "metadata_publisher_match": bool(company and pub_line and company == pub_line),
                "has_published_by": bool(pub_line),
                "has_masthead": False,
                "has_canonical": False,
                "has_footer_repeat": False,
            }
            if chosen:
                source["name"] = chosen
                source["type"] = _classify_org_type(chosen)
                source["confidence"] = _calibrate_confidence("docx", signals)
            else:
                if company:
                    source["candidates"].append({"name": company, "type": "organization", "confidence": 0.5, "why": "DOCX company property"})
                if author:
                    source["candidates"].append({"name": author, "type": "individual", "confidence": 0.35, "why": "DOCX author property"})
            return source, document
        except Exception:
            return source, document

    # TXT
    if ext == ".txt":
        try:
            data = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return source, document
        head = data[:2000]
        tail = data[-2000:]
        pub_line = None
        m = re.search(r'(?i)^(Published|Issued)\s+by\s*:\s*(.+)$', head, re.M)
        if m:
            pub_line = m.group(2).strip()
            add_evidence(m.group(0), section="frontmatter", page=None, start=m.start(0), end=m.end(0), why="Front matter statement")
        c = re.search(r'(?i)©\s+([^\n]+)', head)
        if c:
            c_org = c.group(1).strip()
            add_evidence(f"© {c_org}", section="frontmatter", page=None, start=c.start(0), end=c.end(0), why="Copyright holder")
            if not pub_line:
                pub_line = c_org
        if not pub_line:
            c2 = re.search(r'(?i)©\s+([^\n]+)', tail)
            if c2:
                c_org2 = c2.group(1).strip()
                add_evidence(f"© {c_org2}", section="other", page=None, start=len(data)-2000+c2.start(0), end=len(data)-2000+c2.end(0), why="Tail imprint")
                pub_line = pub_line or c_org2
        if pub_line:
            source["name"] = pub_line
            source["type"] = _classify_org_type(pub_line)
            source["confidence"] = 0.7  # text cues are weaker
        else:
            source["candidates"] = []
        return source, document

    # Fallback: unknown
    return source, document


# --- Document Structure Extraction ---

def extract_document_structure(path: Path, doc_type: Optional[Any]) -> Dict[str, Any]:
    """
    Extract section hierarchy, paragraph numbers, and heading structure.
    Returns structure info with detected patterns.
    """
    ext = path.suffix.lower()
    structure = {
        "has_structure": False,
        "chapters": [],
        "sections": [],
        "paragraphs": [],
        "numbering_pattern": None,
        "toc_detected": False,
        "heading_hierarchy": []
    }

    try:
        if ext == ".pdf":
            return _extract_pdf_structure(path)
        elif ext in (".docx", ".doc"):
            return _extract_docx_structure(path)
        elif ext in (".html", ".htm"):
            return _extract_html_structure(path)
        elif ext == ".txt":
            return _extract_text_structure(path)
    except Exception as e:
        # Don't fail ingestion if structure extraction fails
        pass

    return structure


def _extract_pdf_structure(path: Path) -> Dict[str, Any]:
    """Extract structure from PDF via text analysis."""
    structure = {
        "has_structure": False,
        "chapters": [],
        "sections": [],
        "paragraphs": [],
        "numbering_pattern": None,
        "heading_hierarchy": []
    }

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))

        # Sample first 5 pages for structure detection
        sample_text = ""
        for page in reader.pages[:min(5, len(reader.pages))]:
            sample_text += page.extract_text() + "\n"

        # Detect chapters
        chapter_pattern = re.compile(r'(?:CHAPTER|Chapter|CHAPITRE|Chapitre|Ch\.?)\s+(\d+)', re.IGNORECASE)
        for match in chapter_pattern.finditer(sample_text):
            chapter_num = match.group(1)
            if chapter_num not in [c.get("number") for c in structure["chapters"]]:
                structure["chapters"].append({"number": chapter_num, "title": None})

        # Detect sections with hierarchical numbering (e.g., "1.2.3 Title")
        section_pattern = re.compile(r'^(\d+(?:\.\d+){1,3})\s+([A-Z][^\n]{10,100})', re.MULTILINE)
        for match in section_pattern.finditer(sample_text):
            section_num = match.group(1)
            section_title = match.group(2).strip()
            structure["sections"].append({
                "number": section_num,
                "title": section_title,
                "level": section_num.count('.') + 1
            })

        # Detect paragraph numbering (A.2.2.8, 4.06, etc.)
        para_patterns = [
            (r'\b([A-Z]\.\d+(?:\.\d+){1,3})\b', 'alphanumeric'),  # A.2.2.8
            (r'\b(\d+\.\d+(?:\.\d+)*)\s', 'decimal'),  # 4.06, 12.3.4
            (r'(?:Para|Paragraph|¶)\s*(\d+(?:\.\d+)*)', 'labeled'),  # Paragraph 4.06
        ]

        for pattern, pattern_type in para_patterns:
            matches = re.findall(pattern, sample_text)
            if matches and len(matches) >= 3:  # Need multiple instances to confirm pattern
                structure["numbering_pattern"] = pattern_type
                structure["paragraphs"].extend([{"number": m} for m in matches[:20]])  # Limit to 20
                break

        # Set has_structure flag
        structure["has_structure"] = bool(
            structure["chapters"] or structure["sections"] or structure["paragraphs"]
        )

    except Exception as e:
        pass  # Return empty structure on error

    return structure


def _extract_docx_structure(path: Path) -> Dict[str, Any]:
    """Extract structure from DOCX via style analysis."""
    structure = {
        "has_structure": False,
        "chapters": [],
        "sections": [],
        "paragraphs": [],
        "numbering_pattern": None,
        "heading_hierarchy": []
    }

    try:
        import docx

        doc = docx.Document(str(path))

        # Extract heading hierarchy
        for para in doc.paragraphs[:100]:  # Sample first 100 paragraphs
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                text = para.text.strip()

                if text:
                    structure["heading_hierarchy"].append({
                        "level": level,
                        "text": text
                    })

                    # Check if it's a numbered heading
                    section_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)', text)
                    if section_match:
                        structure["sections"].append({
                            "number": section_match.group(1),
                            "title": section_match.group(2),
                            "level": level
                        })

        # Check for list numbering (paragraph numbers)
        for para in doc.paragraphs[:200]:
            if para.text.strip():
                # Check for alphanumeric paragraph numbers
                para_match = re.match(r'^([A-Z]\.\d+(?:\.\d+)*)\s', para.text)
                if para_match:
                    structure["paragraphs"].append({"number": para_match.group(1)})
                    structure["numbering_pattern"] = "alphanumeric"

        structure["has_structure"] = bool(
            structure["heading_hierarchy"] or structure["sections"] or structure["paragraphs"]
        )

    except Exception as e:
        pass

    return structure


def _extract_html_structure(path: Path) -> Dict[str, Any]:
    """Extract structure from HTML via heading tags."""
    structure = {
        "has_structure": False,
        "chapters": [],
        "sections": [],
        "paragraphs": [],
        "heading_hierarchy": []
    }

    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        # Simple regex-based extraction (could use BeautifulSoup for better parsing)
        heading_pattern = re.compile(r'<h([1-6])[^>]*>([^<]+)</h\1>', re.IGNORECASE)

        for match in heading_pattern.finditer(html_content):
            level = int(match.group(1))
            text = match.group(2).strip()

            structure["heading_hierarchy"].append({
                "level": level,
                "text": text
            })

            # Check for numbered sections
            section_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)', text)
            if section_match:
                structure["sections"].append({
                    "number": section_match.group(1),
                    "title": section_match.group(2),
                    "level": level
                })

        structure["has_structure"] = bool(structure["heading_hierarchy"] or structure["sections"])

    except Exception as e:
        pass

    return structure


def _extract_text_structure(path: Path) -> Dict[str, Any]:
    """Extract structure from plain text via pattern matching."""
    structure = {
        "has_structure": False,
        "sections": [],
        "paragraphs": []
    }

    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        # Look for numbered sections
        section_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+([A-Z][^\n]{10,100})', re.MULTILINE)
        for match in section_pattern.finditer(text):
            structure["sections"].append({
                "number": match.group(1),
                "title": match.group(2).strip(),
                "level": match.group(1).count('.') + 1
            })

        # Look for paragraph numbers
        para_pattern = re.compile(r'\b([A-Z]\.\d+(?:\.\d+)*)\b')
        matches = para_pattern.findall(text)
        if len(matches) >= 3:
            structure["paragraphs"] = [{"number": m} for m in matches[:20]]
            structure["numbering_pattern"] = "alphanumeric"

        structure["has_structure"] = bool(structure["sections"] or structure["paragraphs"])

    except Exception as e:
        pass

    return structure


# --- Progress reporting & result struct ---

@dataclass
class IngestionResult:
    path: Path
    status: str
    duration_s: float
    document_id: Optional[str] = None
    chunks: Optional[int] = None
    error: Optional[str] = None

class ConsoleProgressReporter:
    def __init__(self, quiet: bool = False) -> None:
        self.quiet = quiet
        self._last_progress = -1.0

    def __call__(self, event_type: Any, payload: Optional[Dict[str, Any]] = None) -> None:
        if self.quiet:
            return
        if payload is None and isinstance(event_type, dict):
            payload = event_type
            event_type = payload.get("type")

        event = payload or {}
        etype = str(event_type or event.get("type") or "")
        if etype == "step_progress":
            p = float(event.get("progress", 0.0))
            if p - self._last_progress >= 1.0:
                print(f"   - {event.get('stepId','?')}: {p:.0f}%")
                self._last_progress = p
        elif etype == "step_start":
            print(f"   → {event.get('stepId','?')}: {event.get('message','')}")
        elif etype == "step_complete":
            print(f"   ✓ {event.get('stepId','?')} complete")
        elif etype == "step_error":
            print(f"   ✗ {event.get('stepId','?')} error: {event.get('message','')}")
        elif etype == "error":
            print(f"   ✗ error: {event.get('message','')}")

# --- Environment setup ---

def ensure_embedding_environment(preferred_env_file: Optional[str]) -> None:
    candidates: List[Path] = []
    if preferred_env_file:
        candidates.append(Path(preferred_env_file).expanduser())
    candidates.extend(DEFAULT_ENV_PATHS)
    for c in candidates:
        if c and c.exists():
            try:
                load_env_file(c)
            except Exception as exc:
                print(f"Warning: failed to load env file {c}: {exc}")
    if not os.getenv("OPENAI_API_KEY") and os.getenv("RAG_OPENAI_API_KEY"):
        os.environ["OPENAI_API_KEY"] = os.environ["RAG_OPENAI_API_KEY"]
    if os.getenv("OPENAI_API_KEY"):
        settings.openai_api_key = os.environ["OPENAI_API_KEY"]
    if os.getenv("GOOGLE_API_KEY"):
        settings.google_api_key = os.environ["GOOGLE_API_KEY"]

# --- Ingestion ---

async def ingest_file(
    file_path: Path,
    pipeline: Any,
    base_dir: Path,
    session_id: Optional[str],
    force_refresh: bool,
    tags: List[str],
    source_prefix: Optional[str],
    quiet: bool,
    args: argparse.Namespace,
) -> IngestionResult:
    t0 = time.perf_counter()
    doc_type = infer_document_type(file_path)
    if not doc_type:
        return IngestionResult(path=file_path, status="skipped", duration_s=0.0, error="unsupported extension")

    # Build normalized metadata
    try:
        rel_path = str(file_path.relative_to(base_dir)).replace(os.sep, "/")
    except Exception:
        rel_path = file_path.name

    metadata = {
        "relative_path": rel_path,
        "title": resolve_canonical_title(file_path) or make_title(file_path),
        "source": (source_prefix.rstrip("/") + "/" if source_prefix else "") + rel_path,
        "tags": sorted(set(tags)),
        "session_id": session_id,
        "ingested_at": datetime.now(timezone.utc).isoformat(),
    }

    # Extract source identity & document info per contract
    source_info, document_info = extract_source_identity(file_path, doc_type)

    # Extract document structure (chapters, sections, paragraphs)
    structure_info = extract_document_structure(file_path, doc_type)

    metadata.update({
        "source_identity": source_info,
        "document_info": document_info,
        "structure_info": structure_info,
        "ingestion_options": {
            "preserve_tables": True,
            "table_policy": getattr(args, "table_policy", "both"),
            "table_rows_per_chunk": getattr(args, "table_rows_per_chunk", 10),
            "table_max_cols": getattr(args, "table_max_cols", 40),
            "table_fallback": getattr(args, "table_fallback", "layout"),
            "respect_block_boundaries": True,
            "content_format": "markdown",
            "chunk_meta_propagate_source": True,
        },
    })

    request_type = doc_type
    if not hasattr(doc_type, "value") and isinstance(doc_type, str):
        try:
            request_type = DocumentType(doc_type)
        except Exception:
            request_type = doc_type

    ingest_request = DocumentIngestionRequest(
        file_path=str(file_path),
        type=request_type,
        metadata=metadata,
        force_refresh=force_refresh,
    )

    progress = ConsoleProgressReporter(quiet=quiet)

    try:
        result = await pipeline.ingest_document(ingest_request, progress_callback=progress)
        if hasattr(result, "model_dump"):
            result_payload = result.model_dump()
        elif isinstance(result, dict):
            result_payload = result
        else:
            result_payload = {
                "status": getattr(result, "status", "success"),
                "document_id": getattr(result, "document_id", None),
                "chunks_created": getattr(result, "chunks_created", None),
            }
        return IngestionResult(
            path=file_path,
            status=result_payload.get("status", "success"),
            duration_s=time.perf_counter() - t0,
            document_id=result_payload.get("document_id"),
            chunks=result_payload.get("chunks") or result_payload.get("chunks_created"),
        )
    except Exception as e:
        msg = str(e)
        if "All chunks were duplicates" in msg:
            return IngestionResult(path=file_path, status="duplicate", duration_s=time.perf_counter() - t0)
        return IngestionResult(path=file_path, status="error", duration_s=time.perf_counter() - t0, error=msg)

async def run_ingestion(args: argparse.Namespace) -> int:
    ensure_embedding_environment(args.env_file)

    base_dir = Path(args.base_dir).resolve()
    if not base_dir.exists():
        if args.create_missing_dir:
            base_dir.mkdir(parents=True, exist_ok=True)
        else:
            print(f"Base dir does not exist: {base_dir}")
            return 2

    patterns = tuple(args.patterns) if args.patterns else DEFAULT_PATTERNS
    files = iter_files(patterns, base_dir)
    if not files:
        print("No files matched the provided patterns.")
        return 0

    # Clamp slicing
    if args.offset and args.offset > 0:
        files = files[args.offset:]
    if args.max_files is not None and args.max_files >= 0:
        files = files[: args.max_files]

    batch_size = args.batch_size if args.batch_size and args.batch_size > 0 else 8  # Reduced from 64 to prevent OOM
    total_batches = (len(files) + batch_size - 1) // batch_size

    if args.dry_run:
        print(f"Matched {len(files)} files. Batch size: {batch_size}. Total batches: {total_batches}")
        for p in files[: min(len(files), 50)]:
            dt = infer_document_type(p)
            print(f" - {p} [{(dt.value if hasattr(dt,'value') else dt) if dt else 'unknown'}] {human_size(p.stat().st_size)}")
        if len(files) > 50:
            print(f" ... and {len(files)-50} more")
        return 0

    # Logging level
    os.environ.setdefault("LOG_LEVEL", args.log_level or getattr(settings, "log_level", "INFO"))

    ingest_state = load_ingest_state(base_dir)
    state_changed = False

    vector_store_manager = VectorStoreManager()
    cache_service = CacheService()

    # Ensure vector store + embeddings are ready before ingestion starts
    try:
        await vector_store_manager.initialize()
    except Exception as exc:
        print(f"Failed to initialize vector store manager: {exc}")
        return 1

    # Optional preindex (uses more memory)
    if args.preindex:
        try:
            existing_doc_index: Dict[str, str] = {}
            all_docs = vector_store_manager.get_all_documents()
            for d in all_docs:
                rel = d.get("relative_path") or d.get("metadata", {}).get("relative_path")
                doc_id = d.get("document_id") or d.get("id")
                if rel and doc_id:
                    existing_doc_index[rel] = doc_id
            for file_path in files:
                state_key = str(file_path.resolve())
                if state_key in ingest_state:
                    continue
                try:
                    rel_path = str(file_path.relative_to(base_dir)).replace(os.sep, "/")
                except Exception:
                    rel_path = file_path.name
                doc_id = existing_doc_index.get(rel_path)
                if doc_id:
                    signature = compute_file_signature(file_path)
                    ingest_state[state_key] = {"signature": signature, "document_id": doc_id}
                    state_changed = True
            if state_changed:
                save_ingest_state(base_dir, ingest_state)
        except Exception:
            pass
        finally:
            try:
                del existing_doc_index
            except Exception:
                pass
            gc.collect()

    # Pipeline lifecycle: per batch (default) or per run
    pipeline: Optional[Any] = None
    if args.pipeline_scope == "run":
        pipeline = IngestionPipeline(vector_store_manager, cache_service=cache_service)

    failures: List[IngestionResult] = []
    success_count = 0
    skipped_count = 0
    duplicate_count = 0
    error_count = 0
    global_index = 0

    try:
        for batch_number in range(1, total_batches + 1):
            start_idx = (batch_number - 1) * batch_size
            end_idx = min(start_idx + batch_size, len(files))
            batch_files = files[start_idx:end_idx]
            if not batch_files:
                continue
            print(f"\n=== Batch {batch_number}/{total_batches} ({len(batch_files)} files) ===")

            if args.pipeline_scope == "batch":
                pipeline = IngestionPipeline(vector_store_manager, cache_service=cache_service)

            for file_path in batch_files:
                global_index += 1
                size = human_size(file_path.stat().st_size)
                dt = infer_document_type(file_path)
                type_label = dt.value if hasattr(dt, "value") else str(dt) if dt else "unsupported"
                state_key = str(file_path.resolve())
                signature = compute_file_signature(file_path)
                cached = ingest_state.get(state_key)

                if not args.force_refresh and cached and cached.get("signature") == signature:
                    print(f"[{global_index}/{len(files)}] Skipping {file_path} ({size}, {type_label}) - unchanged")
                    skipped_count += 1
                    continue

                print(f"[{global_index}/{len(files)}] Ingesting {file_path} ({size}, {type_label})")
                result = await ingest_file(
                    file_path=file_path,
                    pipeline=pipeline,
                    base_dir=base_dir,
                    session_id=args.session_id,
                    force_refresh=args.force_refresh,
                    tags=args.tags or [],
                    source_prefix=args.source_prefix,
                    quiet=args.quiet,
                    args=args,
                )

                if result.status == "error":
                    print(f" ✗ Failed: {result.error} ({result.duration_s:.2f}s)")
                    error_count += 1
                    failures.append(result)
                elif result.status == "skipped":
                    print(f" • Skipped: {result.error}")
                    skipped_count += 1
                elif result.status == "duplicate":
                    print(" • Skipped (already ingested)")
                    duplicate_count += 1
                else:
                    print(f" ✓ Status: {result.status} | Doc ID: {result.document_id} | Chunks: {result.chunks} | Time: {result.duration_s:.2f}s")
                    success_count += 1

                # Persist state updates once we have a document id/duplicate
                if result.status in {"success", "completed", "exists", "duplicate"}:
                    ingest_state[state_key] = {"signature": signature, "document_id": result.document_id}
                    state_changed = True

            if state_changed:
                save_ingest_state(base_dir, ingest_state)
                state_changed = False

            # Cleanup between batches
            if args.pipeline_scope == "batch" and pipeline is not None:
                try:
                    await pipeline.cleanup()
                except Exception:
                    pass
                pipeline = None
            gc.collect()

    finally:
        # Run-scoped cleanup
        if pipeline is not None:
            try:
                await pipeline.cleanup()
            except Exception:
                pass
        try:
            await cache_service.disconnect()
        except Exception:
            pass
        try:
            vector_store_manager.executor.shutdown(wait=False)  # type: ignore
        except Exception:
            pass

    print("\nIngestion summary:")
    print(f"  Total files processed: {len(files)}")
    print(f"  Successful ingestions: {success_count}")
    print(f"  Skipped files:         {skipped_count} (duplicates: {duplicate_count})")
    print(f"  Failures:              {error_count}")
    if failures and not args.quiet:
        print("\nFailures:")
        for r in failures[:50]:
            print(f"  - {r.path}: {r.error}")
        if len(failures) > 50:
            print(f"  ... and {len(failures)-50} more")

    return 0 if not failures else 2

# --- CLI ---

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser("ingest_sources_cli", formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    p.add_argument("base_dir", help="Base directory for glob patterns")
    p.add_argument("-p", "--patterns", nargs="*", default=DEFAULT_PATTERNS, help="Glob patterns relative to base_dir")
    p.add_argument("--source-prefix", default=None, help="Logical source prefix to prepend to relative paths")
    p.add_argument("--force-refresh", action="store_true", help="Re-ingest even if file signature matches state")
    p.add_argument("--dry-run", action="store_true", help="List matching files and exit")
    p.add_argument("--max-files", type=int, default=None, help="Cap number of files to process")
    p.add_argument("--offset", type=int, default=0, help="Skip first N files before processing")
    p.add_argument("--batch-size", type=int, default=64, help="Files per batch (memory control)")
    p.add_argument("--pipeline-scope", choices=["batch","run"], default="batch", help="Lifecycle of the ingestion pipeline")
    p.add_argument("--session-id", default=None, help="Optional session id to tag documents")
    p.add_argument("--env-file", default=None, help="Load API keys from this env file (honored even if OPENAI_API_KEY is already set)")
    p.add_argument("--tags", nargs="*", default=[], help="Tags to attach to ingested documents")
    p.add_argument("--quiet", action="store_true", help="Reduce console progress output")
    p.add_argument("--log-level", default=getattr(settings, "log_level", "INFO"), help="Log level (DEBUG, INFO, WARNING, ERROR)")
    p.add_argument("--create-missing-dir", action="store_true", help="Create base_dir if it doesn't exist")

    # Table & chunking options
    p.add_argument("--table-policy", choices=["preserve","rowwise","both"], default="both", help="How to handle tables in chunking")
    p.add_argument("--table-rows-per-chunk", type=int, default=10, help="Max table rows per chunk to avoid splitting rows")
    p.add_argument("--table-max-cols", type=int, default=40, help="Max columns before we drop to rowwise only")
    p.add_argument("--table-fallback", choices=["ocr","layout","skip"], default="layout", help="PDF fallback if native table detection fails")

    # Preindex (off by default to save memory)
    p.add_argument("--preindex", action="store_true", help="Preload vector-store doc index to seed state (uses more memory)")

    args = p.parse_args(argv)
    # Clamp negatives
    if args.offset is not None and args.offset < 0:
        args.offset = 0
    if args.max_files is not None and args.max_files < 0:
        args.max_files = None
    if args.batch_size is None or args.batch_size <= 0:
        args.batch_size = 64
    return args

def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)
    try:
        return asyncio.run(run_ingestion(args))
    except KeyboardInterrupt:
        return 130

if __name__ == "__main__":
    raise SystemExit(main())
